"""PDF to Word conversion using PyMuPDF + python-docx.

Preserves font names, text colors, and paragraph spacing. We deliberately
DON'T use `pdf2docx` (the obvious alternative) because it can hang for
many seconds on PDFs with unusual fonts or complex layouts — instead we
walk PyMuPDF's `dict` representation page-by-page, which is both faster
and bounded.

Conversions run under a soft time budget: `_CONVERSION_TIMEOUT_S` (60s by
default, env-overridable). On timeout we raise `ToolTimeoutError` with a
friendly message so the user knows to try a smaller PDF.
"""
from __future__ import annotations

import concurrent.futures
import io
import logging
import os
import time
import uuid
from typing import Any

import fitz
from docx import Document
from docx.shared import Inches, Pt, RGBColor

from ..utils.cleanup import ensure_temp_dir, get_temp_path
from ..utils.exceptions import ToolTimeoutError

logger = logging.getLogger(__name__)

# Conversion time budget. PDFs that exceed this are almost always going to
# OOM the worker if we let them keep going, so we abort early and surface
# a clear "try smaller" error instead of letting Uvicorn time out.
_CONVERSION_TIMEOUT_S = int(os.environ.get("PDF_TO_WORD_TIMEOUT_S", "60"))


def _fitz_color_to_rgb(color_int: int) -> RGBColor | None:
    """Convert a PyMuPDF color integer to docx RGBColor."""
    if color_int is None or color_int == 0:
        return None  # black / default
    try:
        r = (color_int >> 16) & 0xFF
        g = (color_int >> 8) & 0xFF
        b = color_int & 0xFF
        if r == 0 and g == 0 and b == 0:
            return None
        return RGBColor(r, g, b)
    except Exception:
        return None


def pdf_to_word(input_path: str) -> str:
    """Convert `input_path` to a .docx file and return its path.

    Wraps the actual conversion in a thread + soft timeout. If the
    conversion takes longer than `_CONVERSION_TIMEOUT_S` we raise
    :class:`ToolTimeoutError` so the route handler can surface the
    `Conversion timed out — try a smaller file` message.
    """
    ensure_temp_dir()
    started = time.monotonic()
    try:
        input_size = os.path.getsize(input_path)
    except OSError:
        input_size = 0
    logger.info("pdf_to_word: start input_bytes=%d", input_size)

    with concurrent.futures.ThreadPoolExecutor(max_workers=1) as pool:
        future = pool.submit(_convert_with_open, input_path)
        try:
            output_path = future.result(timeout=_CONVERSION_TIMEOUT_S)
        except concurrent.futures.TimeoutError as exc:
            # We can't cancel the worker thread cleanly (PyMuPDF doesn't
            # honour interrupts mid-render), but the ThreadPoolExecutor
            # going out of scope will detach it and the daemon thread
            # will die with the process. Surface a clean error now so
            # the user isn't left waiting on the front end.
            logger.warning(
                "pdf_to_word: timeout after %ds input_bytes=%d",
                _CONVERSION_TIMEOUT_S, input_size,
            )
            raise ToolTimeoutError(
                f"Conversion took longer than {_CONVERSION_TIMEOUT_S}s — try a smaller PDF"
            ) from exc

    duration_ms = int((time.monotonic() - started) * 1000)
    try:
        out_size = os.path.getsize(output_path)
    except OSError:
        out_size = 0
    logger.info(
        "pdf_to_word: done duration_ms=%d input_bytes=%d output_bytes=%d",
        duration_ms, input_size, out_size,
    )
    return output_path


def _convert_with_open(input_path: str) -> str:
    """Open the PDF and delegate to `_convert_to_docx` with try/finally cleanup."""
    doc = fitz.open(input_path)  # fitz already mmaps the file
    try:
        return _convert_to_docx(doc)
    finally:
        doc.close()


def _extract_page_payload(page) -> dict:
    """Extract the structured `dict` payload for a single page.

    Runs in a thread pool — PyMuPDF page methods release the GIL on the
    heavy C-side work (text shaping, image decode), so parallelising
    helps multi-page documents finish substantially faster on multi-core
    boxes. We do NOT touch the docx writer here — python-docx isn't
    thread-safe for shared-Document writes.
    """
    return {
        "text": page.get_text("text") or "",
        "blocks": page.get_text("dict").get("blocks", []),
    }


def _convert_to_docx(doc) -> str:
    """Render `doc` to a .docx file and return its path."""
    # Bail out early on image-only PDFs — the rest of the pipeline produces
    # a blank .docx in that case, which silently confuses users. The error
    # string contains "OCR" + "no text" so the frontend friendlyError() can
    # route it to a clear UX message.
    has_any_text = False
    for page in doc:
        if (page.get_text("text") or "").strip():
            has_any_text = True
            break
    if not has_any_text:
        raise ValueError(
            "This PDF has no text layer — run OCR PDF first to make it searchable"
        )

    word = Document()

    # Set default style
    style = word.styles["Normal"]
    style.font.size = Pt(11)

    # Pre-extract every page's block dict in parallel. PyMuPDF releases
    # the GIL for the heavy lifting (text shaping, image decode) so this
    # genuinely speeds up multi-page documents on multi-core hosts.
    page_count = len(doc)
    pages_payload: list[dict[str, Any]] = [None] * page_count  # type: ignore[list-item]
    if page_count > 1:
        max_workers = max(1, (os.cpu_count() or 2) - 1)
        max_workers = min(max_workers, 4, page_count)
        with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as pool:
            futures = {
                pool.submit(_extract_page_payload, doc[i]): i for i in range(page_count)
            }
            for fut in concurrent.futures.as_completed(futures):
                idx = futures[fut]
                pages_payload[idx] = fut.result()
    else:
        pages_payload[0] = _extract_page_payload(doc[0])

    for page_num in range(page_count):
        if page_num > 0:
            word.add_page_break()

        blocks = pages_payload[page_num]["blocks"]
        for block in blocks:
            if block.get("type") == 0:  # text block
                for line in block.get("lines", []):
                    para = word.add_paragraph()
                    # Reduce paragraph spacing for tighter layout
                    para.paragraph_format.space_before = Pt(0)
                    para.paragraph_format.space_after = Pt(1)

                    for span in line.get("spans", []):
                        run = para.add_run(span["text"])

                        # Font size with clamping
                        font_size = max(6, min(span.get("size", 11), 72))
                        run.font.size = Pt(font_size)

                        # Preserve font name
                        font_name = span.get("font", "")
                        if font_name:
                            # Map common PDF font names to Word-compatible names
                            name_lower = font_name.lower()
                            if "arial" in name_lower or "helvetica" in name_lower:
                                run.font.name = "Arial"
                            elif "times" in name_lower:
                                run.font.name = "Times New Roman"
                            elif "courier" in name_lower:
                                run.font.name = "Courier New"
                            elif "calibri" in name_lower:
                                run.font.name = "Calibri"
                            else:
                                run.font.name = font_name.split("+")[-1].split("-")[0]

                        # Preserve text color
                        color_val = span.get("color")
                        if color_val and color_val != 0:
                            rgb = _fitz_color_to_rgb(color_val)
                            if rgb:
                                run.font.color.rgb = rgb

                        # Font flags
                        flags = span.get("flags", 0)
                        if flags & (1 << 0):  # superscript
                            run.font.superscript = True
                        if flags & (1 << 4):  # bold
                            run.font.bold = True
                        if flags & (1 << 1):  # italic
                            run.font.italic = True

            elif block.get("type") == 1:  # image block
                try:
                    img_data = block.get("image")
                    if img_data:
                        img_stream = io.BytesIO(img_data)
                        word.add_picture(img_stream, width=Inches(5))
                except Exception:
                    pass

    output_path = get_temp_path(f"converted_{uuid.uuid4().hex}.docx")
    word.save(str(output_path))
    return str(output_path)
